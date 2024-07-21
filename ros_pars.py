import sys
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QFileDialog, QMessageBox
import docx
import requests
from docx import Document

dict_of_categoty_type = {'003001000000': 'Земли селькохозяйственного назначения',
                         '003003000000': 'Земли промышленности, энергетики, транспорта, связи, радиовещания, '
                                         'телевидения, информатики, земли для обеспечения космической деятельности, '
                                         'земли обороны, безопасности и земли иного специального назначения',
                         '003008000000': 'Категория не установлена',
                         None: '-'}

dict_of_type_of_own = {100: 'Частная собственность',
                       200: 'Собственность публично-правовых образований',
                       None: '-'}


class AppWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.left = 500
        self.top = 500
        self.setWindowTitle("self-recorder")
        self.first_table = None  # Инициализация переменной для первой таблицы
        self.second_table = None  # Инициализация переменной для второй таблицы

        layout = QVBoxLayout()

        btn1 = QPushButton("Исходная таблица")
        btn2 = QPushButton("Конечная таблица")
        btn3 = QPushButton("Выполнить")

        btn1.clicked.connect(self.open_first_table)
        btn2.clicked.connect(self.open_second_table)
        btn3.clicked.connect(self.main_bt)

        layout.addWidget(btn1)
        layout.addWidget(btn2)
        layout.addWidget(btn3)

        self.list_of_address = []
        self.list_of_category_type = []
        self.list_of_allows_to_use = []
        self.list_of_type_of_own = []
        self.setLayout(layout)

    def transform_string(self, input_string):
        parts = input_string.split(':')
        for i in range(len(parts)):
            if i == 2:  # обрабатываем третий элемент
                parts[i] = str(int(parts[i]))  # убираем ведущие нули
            else:
                parts[i] = str(int(parts[i])) if int(parts[i]) != 0 else parts[
                    i]  # убираем ведущие нули, но оставляем 0
        result = ':'.join(parts)
        return result

    def read_first_column_from_docx(self, file_path):
        doc = docx.Document(file_path)
        table = doc.tables[0]  # Предполагается, что таблица находится на первой странице
        column_data = []
        for row in table.rows:
            cell = row.cells[0]  # Получаем только ячейку из первого столбца
            column_data.append(cell.text)
        return column_data

    def open_first_table(self):
        options = QFileDialog.Options()
        self.first_table, _ = QFileDialog.getOpenFileName(self, "Выберите файл", "",
                                                   "Word Documents (*.docx);;All Files (*);;Python Files (*.py)",
                                                   options=options)
        if self.first_table:  # Убеждаемся, что пользователь выбрал файл
            table_data = self.read_first_column_from_docx(self.first_table)
            unique_elements_set = set(item.strip() for item in table_data)
            self.main_unique_elements_list = list(unique_elements_set)
            self.unique_elements_list_for_request = []
            for i in self.main_unique_elements_list:
                self.unique_elements_list_for_request.append(self.transform_string(i))

            print("Выбранный исходный файл:",
                  self.first_table)  # пока нужен для отслеживания действия в консоли

    def open_second_table(self):
        options = QFileDialog.Options()
        self.second_table, _ = QFileDialog.getOpenFileName(self, "Выберите файл", "",
                                                           "Word Documents (*.docx);;All Files (*);;Python Files (*.py)",
                                                           options=options)
        if self.second_table:
            print("Выбран файл для второй таблицы:", self.second_table)
        return self.second_table

    def main_bt(self):
        if self.first_table and self.second_table:
            count = 1
            for number in self.unique_elements_list_for_request:
                url = r'https://pkk.rosreestr.ru/api/features/1/' + str(number)
                response = requests.get(url, verify=False) # игнорирую верификацию т.к. проблема с сертификатом
                if response.status_code == 200:
                    print(f'обработано {count} участков из {len(self.main_unique_elements_list)}')
                    count += 1
                    data = response.json()
                    if data['feature']['attrs']['address'] is not None:
                        self.list_of_address.append(data['feature']['attrs']['address'])
                    else:
                        self.list_of_address.append('-')
                    if data['feature']['attrs']['util_by_doc'] is not None:
                        self.list_of_allows_to_use.append(data['feature']['attrs']['util_by_doc'])
                    else:
                        self.list_of_allows_to_use.append('-')
                    self.list_of_category_type.append(dict_of_categoty_type.get(data['feature']['attrs']['category_type'],
                                                                       '███████Данного значения еще нет в БД,внесите изменения███████'))
                    self.list_of_type_of_own.append(dict_of_type_of_own.get(data['feature']['attrs']['fp'],
                                                                   '███████Данного значения еще нет в БД,внесите изменения████████'))
                else:
                    self.list_of_address.append('█████Request error:', response.status_code)
                    self.list_of_category_type.append('█████Request error:', response.status_code)
                    self.list_of_allows_to_use.append('█████Request error:', response.status_code)
                    self.ist_of_type_of_own.append('█████Request error:', response.status_code)

            # Код, относящийся к таблице doc, перемещен внутрь условия
            doc = Document(self.second_table)
            table = doc.tables[0]
            for i in range(len(self.main_unique_elements_list)):
                table.cell(i + 2, 0).text = self.main_unique_elements_list[i]
                table.cell(i + 2, 1).text = self.list_of_address[i]
                table.cell(i + 2, 2).text = self.list_of_category_type[i]
                table.cell(i + 2, 3).text = self.list_of_allows_to_use[i]
                table.cell(i + 2, 4).text = self.list_of_type_of_own[i]
            doc.save(self.second_table)
            QMessageBox.information(self, 'Отчет', 'готово')
        else:
            QMessageBox.information(self, 'Предупреждение', 'Нужно указать путь к обеим таблицам')


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = AppWindow()
    window.show()
    sys.exit(app.exec_())
